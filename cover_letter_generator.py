from docx import Document
import openpyxl
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import os

# Initialize global variables
output_directory = 'New Letter of Credits'


def process_excel_files(excel_file_paths):
    if not word_file_path:
        messagebox.showwarning("No Word File", "Please select a Word document.")
        return

    # Create the output directory if it doesn't exist
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    for excel_file_path in excel_file_paths:
        workbook = openpyxl.load_workbook(excel_file_path, data_only=True)
        sheet = workbook['CI']

        Date = datetime.now().strftime('%Y.%m.%d')
        to, units, lc_number, invoice_number, invoice_value, payable_value = "", "", "", "", "", ""

        for row in sheet.iter_rows(values_only=True):
            for idx, cell in enumerate(row):
                if cell == "INVOICE NO.:":
                    invoice_number = row[idx + 1]
                elif cell == "TO:":
                    to = row[idx + 1].split("\n")[0]
                elif cell == "TOTAL":
                    units = row[idx + 3]
                elif cell is None and idx + 2 < len(row) and row[idx + 1] == "EUR" and row[idx - 1] is None:
                    invoice_value = row[idx + 2]
                elif (cell == "TOTAL AMOUNT" or cell == "TOTAL CIF") and row[idx + 1] == "EUR":
                    payable_value = row[idx + 2]
                elif isinstance(cell, str) and cell.startswith("THE DOCUMENTARY CREDIT NUMBER:"):
                    parts = cell.split(":")
                    lc_number = parts[1].split("\n")[0]

        # Format the invoice values
        try:
            in_value = float(invoice_value)
            pay_value = float(payable_value)
            formatted_invoice_value = f"{in_value:,.2f}"
            formatted_payable_value = f"{pay_value:,.2f}"
        except ValueError:
            formatted_invoice_value = invoice_value
            formatted_payable_value = payable_value

        # Load the Word document
        doc = Document(word_file_path)

        for para in doc.paragraphs:
            if "Date:" in para.text:
                para.text = para.text.replace("2024.01.01", Date)
            if "company, units1, LC number: x" in para.text:
                para.clear()
                replacement_text = [to, f"{units}-units", f"LC number: {lc_number}"]
                for i, text in enumerate(replacement_text):
                    run = para.add_run(text)
                    run.bold = True
                    if i < len(replacement_text) - 1:
                        para.add_run(", ")
            if "units2" in para.text:
                para.text = para.text.replace("units2", f"{units}-units")
            if "invoice number:" in para.text:
                para.text = para.text.replace(
                    "invoice number: x, invoice value:  x EUR, 95% of payable invoice value: x EUR",
                    f"invoice number: {invoice_number}, invoice value: {formatted_invoice_value} EUR, 95% of payable invoice value: {formatted_payable_value} EUR"
                )

        # Save the modified document
        output_file_name = f'Letter of Credits to {to} {invoice_number}.docx'
        save_path = os.path.join(output_directory, output_file_name)

        doc.save(save_path)

    messagebox.showinfo("Success", "Word documents saved successfully.")

# GUI Setup
def select_excel_files():
    files = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx")])
    if files:
        file_listbox.delete(0, tk.END)
        for file in files:
            file_listbox.insert(tk.END, file)

def select_word_file():
    global word_file_path
    word_file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])

def process_files():
    excel_files = file_listbox.get(0, tk.END)
    if not excel_files:
        messagebox.showwarning("No Files", "Please select at least one Excel file.")
    elif not word_file_path:
        messagebox.showwarning("No Word File", "Please select a Word document.")
    else:
        process_excel_files(excel_files)

root = tk.Tk()
root.title("Letter of Credit Generator")
root.geometry("500x400")

file_listbox = tk.Listbox(root, selectmode=tk.MULTIPLE, width=60, height=15)
file_listbox.pack(pady=10, padx=10, fill='both', expand=True)

scrollbar = tk.Scrollbar(file_listbox)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

file_listbox.config(yscrollcommand=scrollbar.set)
scrollbar.config(command=file_listbox.yview)

tk.Button(root, text="Select Excel Files", command=select_excel_files).pack(pady=5)
tk.Button(root, text="Select Word Document", command=select_word_file).pack(pady=5)
tk.Button(root, text="Process and Save", command=process_files).pack(pady=20)

root.mainloop()
